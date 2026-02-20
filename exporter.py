from fpdf import FPDF
from docx import Document
from models import Exam
import json
import io
import os
import tempfile

def export_to_pdf(exam: Exam, filename: str = None):
    pdf = FPDF()
    pdf.add_page()
    
    pdf.add_page()
    
    try:
        pdf.add_font('Arial', '', 'C:\\Windows\\Fonts\\arial.ttf', uni=True)
        pdf.set_font("Arial", size=12)
    except:
        pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt=exam.title, ln=1, align='C')
    pdf.ln(10)
    
    for i, q in enumerate(exam.questions, 1):
        question_text = f"{i}. {q.text} ({q.difficulty})"
        pdf.multi_cell(0, 10, txt=question_text)
            
        if q.q_type == 'Multiple Choice' and q.options:
            options = json.loads(q.options)
            for opt in options:
                pdf.cell(10) # Indent
                pdf.cell(0, 10, txt=f"- {opt}", ln=1)
        pdf.ln(5)
        
    # --- ANSWER KEY ---
    pdf.add_page()
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(0, 10, txt="CEVAP ANAHTARI", ln=1, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=12)
    
    
    for i, q in enumerate(exam.questions, 1):
        pdf.cell(0, 8, txt=f"{i}. {q.answer}", ln=1)

    if filename:
        pdf.output(filename)
        print(f"PDF exported to {filename}")
        return None
    else:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf.output(tmp.name)
            tmp.close()
            with open(tmp.name, "rb") as f:
                buffer = io.BytesIO(f.read())
            try:
                os.remove(tmp.name)
            except:
                pass
            return buffer

def export_to_docx(exam: Exam, filename: str = None):
    doc = Document()
    doc.add_heading(exam.title, 0)
    
    for i, q in enumerate(exam.questions, 1):
        doc.add_paragraph(f"{i}. {q.text} ({q.difficulty})")
        
        if q.q_type == 'Multiple Choice' and q.options:
            options = json.loads(q.options)
            for opt in options:
                doc.add_paragraph(f"- {opt}", style='List Bullet')
    
    # --- ANSWER KEY ---
    doc.add_page_break()
    doc.add_heading("CEVAP ANAHTARI", level=1)
    
    for i, q in enumerate(exam.questions, 1):
        doc.add_paragraph(f"{i}. {q.answer}")
                
    if filename:
        doc.save(filename)
        print(f"Word doc exported to {filename}")
        return None
    else:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
